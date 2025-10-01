"""DOCX document extractor."""

from __future__ import annotations

import io
import logging
import re
import time
from pathlib import Path
from typing import Any, ClassVar

from docx import Document
from files.backend.extractors.base import DocumentExtractor, ExtractionResult
from PIL import Image

logger = logging.getLogger(__name__)


class DOCXExtractor(DocumentExtractor):
    """Extract content from DOCX documents"""

    SUPPORTED_EXTENSIONS: ClassVar[set[str]] = {".docx", ".doc"}

    async def can_extract(self, file_path: Path) -> bool:
        """Check if this extractor can handle the given file"""
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    async def extract(self, file_path: Path, options: dict[str, Any] | None = None) -> ExtractionResult:
        """Extract content from DOCX"""
        start_time = time.time()
        options = options or {}

        # Validate file
        await self.validate_file(file_path)

        result = ExtractionResult(
            file_path=str(file_path),
            file_type="docx",
            file_size=file_path.stat().st_size,
            extraction_method="python-docx",
            processing_time_ms=0,
        )

        try:
            doc = Document(str(file_path))

            self._populate_metadata(result, doc)
            result.sections = self._build_sections(doc)
            result.full_text = self._collect_full_text(doc)

            if options.get("extract_tables", True):
                result.tables = await self._extract_tables(doc)

            if options.get("extract_images", False):
                result.images = await self._extract_images(doc)

        except Exception as e:
            logger.exception(f"Error extracting DOCX {file_path}")
            result.error_messages.append(str(e))

        result.processing_time_ms = int((time.time() - start_time) * 1000)
        return result

    def _get_heading_level_from_style(self, style_name: str) -> int:
        """Extract heading level from style name"""
        if "Heading" not in style_name:
            return 1

        # Try to extract number from style name

        match = re.search(r"\d+", style_name)
        if match:
            level = int(match.group())
            return min(level, 6)

        return 1

    def _populate_metadata(self, result: ExtractionResult, doc: Any) -> None:
        """Populate basic metadata from document core properties."""
        core = doc.core_properties
        if not core:
            return

        result.title = core.title
        result.author = core.author
        result.subject = core.subject
        if core.keywords:
            result.keywords = [keyword.strip() for keyword in core.keywords.split(",")]
        result.language = core.language or "en"

    def _build_sections(self, doc: Any) -> list[dict[str, Any]]:
        """Convert paragraphs into structured sections."""
        sections: list[dict[str, Any]] = []
        current_section: dict[str, Any] | None = None
        current_content: list[str] = []
        section_counter = 0

        for paragraph in doc.paragraphs:
            if self._is_heading_paragraph(paragraph):
                current_section = self._start_new_section(paragraph, sections, current_section, current_content, section_counter)
                section_counter += 1
                current_content = []
                continue

            if paragraph.text.strip():
                current_content.append(paragraph.text)

        self._flush_section(sections, current_section, current_content)
        return sections or self._build_default_section(doc)

    def _is_heading_paragraph(self, paragraph: Any) -> bool:
        """Return True when the paragraph represents a heading."""
        return bool(paragraph.style and "Heading" in paragraph.style.name)

    def _start_new_section(
        self,
        paragraph: Any,
        sections: list[dict[str, Any]],
        current_section: dict[str, Any] | None,
        current_content: list[str],
        section_counter: int,
    ) -> dict[str, Any]:
        """Close the current section and create the next one."""
        self._flush_section(sections, current_section, current_content)

        next_index = section_counter + 1
        level = self._get_heading_level_from_style(paragraph.style.name)
        return {
            "level": level,
            "title": paragraph.text.strip(),
            "order": next_index,
            "path": str(next_index),
            "style": paragraph.style.name,
        }

    def _flush_section(
        self,
        sections: list[dict[str, Any]],
        current_section: dict[str, Any] | None,
        current_content: list[str],
    ) -> None:
        """Finalize and append the current section if it has content."""
        if not current_section:
            return

        current_section = {**current_section}  # Avoid mutating shared dict state
        current_section["content"] = "\n".join(current_content).strip()
        sections.append(current_section)

    def _build_default_section(self, doc: Any) -> list[dict[str, Any]]:
        """Create a single section when no structured headings are present."""
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        if not text:
            return []
        return [
            {
                "level": 1,
                "title": "Document Content",
                "content": text,
                "order": 1,
                "path": "1",
            }
        ]

    def _collect_full_text(self, doc: Any) -> str:
        """Concatenate entire document text."""
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    async def _extract_tables(self, doc: Any) -> list[dict[str, Any]]:
        """Extract tables from document"""
        tables: list[dict[str, Any]] = []

        try:
            for index, table in enumerate(doc.tables):
                if not table.rows:
                    continue

                headers = self._extract_table_headers(table)
                rows = self._extract_table_rows(table, headers)
                if not rows:
                    continue

                tables.append(
                    {
                        "name": f"Table_{index}",
                        "headers": headers,
                        "rows": rows,
                        "schema": self.infer_table_schema(headers, rows),
                    }
                )

        except Exception as exc:  # pragma: no cover - defensive for third-party parser
            logger.warning(f"Failed to extract tables: {exc}")

        return tables

    def _extract_table_headers(self, table: Any) -> list[str]:
        """Return normalized table headers."""
        headers: list[str] = []
        for cell in table.rows[0].cells:
            header_text = cell.text.strip() or f"Column_{len(headers)}"
            headers.append(header_text)
        return headers

    def _extract_table_rows(self, table: Any, headers: list[str]) -> list[dict[str, Any]]:
        """Return structured data rows for a table."""
        rows: list[dict[str, Any]] = []
        for row in table.rows[1:]:
            row_data = {headers[idx]: cell.text.strip() for idx, cell in enumerate(row.cells) if idx < len(headers)}
            if any(row_data.values()):
                rows.append(row_data)
        return rows

    async def _extract_images(self, doc: Any) -> list[dict[str, Any]]:
        """Extract images from document"""
        images = []

        try:
            # Access document relationships
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_part = rel.target_part

                    # Get image data
                    image_data = image_part.blob

                    # Determine image type from content type
                    content_type = image_part.content_type

                    # Get image filename
                    image_name = image_part.partname.split("/")[-1]

                    image_info = {
                        "name": image_name,
                        "mime_type": content_type,
                        "file_size": len(image_data),
                    }

                    # Try to get dimensions if PIL is available
                    try:
                        img = Image.open(io.BytesIO(image_data))
                        image_info["dimensions"] = {
                            "width": img.width,
                            "height": img.height,
                        }
                    except ImportError:
                        pass
                    except Exception as e:
                        logger.debug(f"Could not get image dimensions: {e}")

                    images.append(image_info)

        except Exception as e:
            logger.warning(f"Failed to extract images: {e}")

        return images
